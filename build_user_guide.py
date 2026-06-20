from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
ASSETS = ROOT / "guide-assets"
ASSETS.mkdir(exist_ok=True)

NAVY = "#223A66"
BLUE = "#6E96CF"
PALE = "#E8F1FB"
MINT = "#BFE6D5"
INK = "#1C2D4D"
MUTED = "#71809B"
WHITE = "#FFFFFF"
GOLD = "#E6BE67"


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def save_setup_illustration():
    im = Image.new("RGB", (1400, 820), "#F5F8FC")
    d = ImageDraw.Draw(im)
    rounded(d, (40, 35, 1360, 785), 28, WHITE, "#DCE5F1", 2)
    d.text((75, 68), "EVENT BUILDER", font=font(21, True), fill=BLUE)
    d.text((75, 105), "Choose your game mode", font=font(43, True), fill=INK)
    modes = [
        ("OPEN PLAY", "Fair individual rotation", "#E6EFFB"),
        ("SKILL LADDER", "Match similar results", "#EEF6F2"),
        ("TEAM TOURNAMENT", "Fixed-team competition", "#FFF5DE"),
    ]
    for i, (title, sub, color) in enumerate(modes):
        x = 75 + i * 285
        rounded(d, (x, 180, x + 255, 325), 18, color, BLUE if i == 0 else "#DCE5F1", 3 if i == 0 else 1)
        d.text((x + 18, 205), title, font=font(18, True), fill=INK)
        d.text((x + 18, 250), sub, font=font(14), fill=MUTED)
    rounded(d, (75, 360, 885, 720), 18, "#FAFCFE", "#DCE5F1")
    fields = [
        ("Session name", "Saturday Open Play"),
        ("Courts", "4"),
        ("Game format", "Doubles"),
        ("Play to / Win by", "11 / 2"),
        ("Scoring", "Traditional side-out"),
        ("Matchmaking", "Balanced remix"),
    ]
    for i, (label, value) in enumerate(fields):
        col = i % 2
        row = i // 2
        x = 105 + col * 375
        y = 395 + row * 95
        d.text((x, y), label.upper(), font=font(12, True), fill=MUTED)
        rounded(d, (x, y + 25, x + 335, y + 70), 10, WHITE, "#C9D7E8")
        d.text((x + 13, y + 38), value, font=font(16, True), fill=INK)
    rounded(d, (920, 180, 1315, 720), 22, NAVY)
    d.text((950, 210), "EVENT SIMULATION", font=font(14, True), fill="#BCD4F0")
    d.text((950, 250), "Fair Rotation", font=font(30, True), fill=WHITE)
    rounded(d, (950, 310, 1285, 500), 15, "#6E96CF")
    for x, y, name in [(985, 345, "A1"), (1135, 345, "B1"), (985, 435, "A2"), (1135, 435, "B2")]:
        d.text((x, y), name, font=font(20, True), fill=WHITE)
    d.text((1092, 395), "VS", font=font(29, True), fill="#FFF1C6")
    for i, line in enumerate(["1  Register players", "2  Fill open courts", "3  Record results", "4  Rotate fairly"]):
        d.text((950, 535 + i * 38), line, font=font(15), fill="#D8E5F0")
    path = ASSETS / "event-builder.png"
    im.save(path)
    return path


def save_court_illustration():
    im = Image.new("RGB", (1400, 820), "#F5F8FC")
    d = ImageDraw.Draw(im)
    rounded(d, (50, 40, 1350, 780), 26, WHITE, "#DCE5F1", 2)
    d.text((85, 75), "COURT 1", font=font(26, True), fill=INK)
    d.text((1160, 82), "PLAYING", font=font(15, True), fill="#4C8C77")
    rounded(d, (85, 135, 1315, 330), 20, "#EDF3FA", "#D8E3F0")
    rounded(d, (85, 135, 700, 330), 20, "#E2EDFA")
    rounded(d, (700, 135, 1315, 330), 20, "#E8F2EE")
    d.text((120, 175), "TEAM A", font=font(18, True), fill=BLUE)
    d.text((120, 225), "Angela Ocampo", font=font(28, True), fill=INK)
    d.text((120, 270), "Marco Valdez", font=font(28, True), fill=INK)
    d.text((1130, 175), "TEAM B", font=font(18, True), fill="#6D9F90")
    d.text((1000, 225), "Ysa Torres", font=font(28, True), fill=INK)
    d.text((1010, 270), "Ken Ramos", font=font(28, True), fill=INK)
    d.ellipse((655, 207, 745, 297), fill=NAVY, outline=WHITE, width=7)
    d.text((677, 231), "VS", font=font(25, True), fill=WHITE)
    d.text((85, 370), "RECORD THE GAME", font=font(15, True), fill=MUTED)
    controls = [
        ("Team A won rally", "#EDF4FC"),
        ("Team B won rally", "#EDF4FC"),
        ("Team A won game", "#F5FAF8"),
        ("Team B won game", "#F5FAF8"),
    ]
    for i, (label, color) in enumerate(controls):
        x = 85 + (i % 2) * 310
        y = 405 + (i // 2) * 70
        rounded(d, (x, y, x + 285, y + 52), 11, color, "#CAD9E9")
        d.text((x + 20, y + 16), label, font=font(15, True), fill=INK)
    rounded(d, (740, 370, 1315, 605), 17, "#FAFCFE", "#DCE5F1")
    d.text((775, 400), "FINAL SCORE", font=font(15, True), fill=MUTED)
    rounded(d, (775, 445, 910, 515), 10, WHITE, "#C9D7E8")
    rounded(d, (980, 445, 1115, 515), 10, WHITE, "#C9D7E8")
    d.text((820, 463), "11", font=font(28, True), fill=INK)
    d.text((1025, 463), "9", font=font(28, True), fill=INK)
    rounded(d, (1140, 445, 1285, 515), 10, BLUE)
    d.text((1161, 468), "SAVE", font=font(17, True), fill=WHITE)
    d.text((775, 550), "Substitute player  •  Close court", font=font(16, True), fill=MUTED)
    rounded(d, (85, 650, 1315, 725), 13, NAVY)
    d.text((120, 672), "NEXT: rested players with fewer games are assigned automatically", font=font(19, True), fill=WHITE)
    path = ASSETS / "court-control.png"
    im.save(path)
    return path


def save_tournament_illustration():
    im = Image.new("RGB", (1400, 820), "#F5F8FC")
    d = ImageDraw.Draw(im)
    d.text((55, 45), "TEAM TOURNAMENT", font=font(22, True), fill=BLUE)
    d.text((55, 82), "Standings & schedule", font=font(42, True), fill=INK)
    rounded(d, (55, 150, 1345, 500), 20, WHITE, "#DCE5F1")
    headers = ["#", "TEAM", "PLAYED", "WON", "LOST", "PF", "DIFF"]
    xs = [85, 160, 650, 800, 930, 1060, 1190]
    rounded(d, (55, 150, 1345, 210), 20, NAVY)
    for x, h in zip(xs, headers):
        d.text((x, 170), h, font=font(14, True), fill=WHITE)
    rows = [
        ("1", "Kitchen Crushers", "4", "4", "0", "44", "+21"),
        ("2", "Dink Dynasty", "4", "3", "1", "39", "+12"),
        ("3", "Paddle Patrol", "4", "2", "2", "35", "+3"),
        ("4", "Drop Shot Duo", "4", "1", "3", "27", "-11"),
    ]
    for r, values in enumerate(rows):
        y = 235 + r * 62
        if r % 2:
            d.rectangle((56, y - 12, 1344, y + 43), fill="#F5F8FC")
        for x, value in zip(xs, values):
            d.text((x, y), value, font=font(17, r == 0), fill=INK)
    d.text((55, 545), "UPCOMING MATCHES", font=font(16, True), fill=MUTED)
    matches = [("R1", "Spin Doctors", "vs", "Court Commanders"), ("R2", "Kitchen Crushers", "vs", "Paddle Patrol")]
    for i, (rnd, a, vs, b) in enumerate(matches):
        x = 55 + i * 650
        rounded(d, (x, 580, x + 610, 705), 17, WHITE, "#DCE5F1")
        rounded(d, (x + 20, 610, x + 72, 660), 10, PALE)
        d.text((x + 34, 626), rnd, font=font(14, True), fill=BLUE)
        d.text((x + 100, 615), a, font=font(18, True), fill=INK)
        d.text((x + 280, 615), vs, font=font(15, True), fill=MUTED)
        d.text((x + 330, 615), b, font=font(18, True), fill=INK)
        d.text((x + 100, 650), "Assigned automatically when a court opens", font=font(13), fill=MUTED)
    path = ASSETS / "tournament.png"
    im.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def build_docx():
    setup_img = save_setup_illustration()
    court_img = save_court_illustration()
    tourney_img = save_tournament_illustration()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1C2D4D")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 30, "223A66", 0, 8),
        ("Heading 1", 18, "476FAE", 16, 8),
        ("Heading 2", 14, "476FAE", 12, 6),
        ("Heading 3", 12, "223A66", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    header = sec.header.paragraphs[0]
    header.text = "KING OF OPEN PLAY  |  ORGANIZER USER GUIDE"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("71809B")
    add_page_number(sec.footer.paragraphs[0])

    logo = doc.add_picture(str(ROOT / "openplay-ph-logo.png"), width=Inches(1.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("King of Open Play")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Organizer User Guide")
    run.font.name = "Arial"
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("6E96CF")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Open Play • Skill Ladder • Team Tournament • Offline Scoring")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("71809B")
    doc.add_paragraph()
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = lead.add_run("A practical field manual for setting up the event, managing courts, recording results, and protecting session data.")
    rr.font.size = Pt(12)
    rr.font.italic = True
    rr.font.color.rgb = RGBColor.from_string("476FAE")
    doc.add_page_break()

    doc.add_heading("Quick Start", level=1)
    steps = [
        ("Build the event", "Choose Open Play, Skill Ladder, or Team Tournament. Configure courts, scoring, and matchmaking."),
        ("Register participants", "Add individual player names or fixed tournament teams."),
        ("Fill courts", "Let the app choose eligible players or scheduled teams."),
        ("Record results", "Track every rally, tap the game winner, or enter the final score."),
        ("Keep the rotation moving", "The next eligible match is assigned automatically."),
        ("Export before clearing data", "Download a JSON backup for important sessions."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    for i, (label, detail) in enumerate(steps, 1):
        cells = table.add_row().cells
        cells[0].width = Inches(1.75)
        cells[1].width = Inches(4.75)
        cells[0].text = f"{i}. {label}"
        cells[1].text = detail
        set_cell_shading(cells[0], "E6EFFB")
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("223A66")
    table.rows[0]._element.getparent().remove(table.rows[0]._element)

    doc.add_heading("1. Build the Event", level=1)
    doc.add_paragraph("Select Build an event, New session, or the settings icon. The Event Builder shows what each mode will do before the organizer starts.")
    doc.add_picture(str(setup_img), width=Inches(6.45))
    cap = doc.add_paragraph("Figure 1. Event Builder and gameplay simulation.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor.from_string("71809B")
    for heading, text in [
        ("Open Play", "Best for social sessions. Players rotate individually and equal playing time is prioritized."),
        ("Skill Ladder", "Best for competitive open play. Results gradually improve future skill matching."),
        ("Team Tournament", "Best for organized competition. Partners remain fixed and standings determine rank."),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{heading}: ").bold = True
        p.add_run(text)

    doc.add_heading("2. Register Players or Teams", level=1)
    doc.add_paragraph("In Open Play and Skill Ladder, open Players and enter one name per line. The bulk entry supports large groups. Duplicate names are ignored.")
    doc.add_paragraph("In Team Tournament, register a team name and two players. Existing individual players can be paired automatically using Convert players into teams.")

    doc.add_heading("3. Run Live Courts", level=1)
    doc.add_picture(str(court_img), width=Inches(6.45))
    cap = doc.add_paragraph("Figure 2. Live court controls and side-by-side matchup.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor.from_string("71809B")
    court_points = [
        "Fill open courts assigns rested players with fewer games first.",
        "The announcer reads the court and player names when enabled.",
        "Rally buttons follow the selected side-out or rally-scoring rules.",
        "Winner-only entry is the fastest result option.",
        "Final-score entry supplies better point-margin skill data.",
        "Substitute a player replaces someone who leaves early without resetting the score.",
        "Close court removes a court from automatic assignments until it is reopened.",
    ]
    for item in court_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Scoring Rules and Validation", level=1)
    doc.add_paragraph("Traditional side-out scoring starts doubles at 0-0-2, awards points only to the serving team, and tracks server one, server two, and side-outs.")
    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Inches(6.4)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "EDF4FC")
    cell.text = "Example — Play To 11, Win By 2: 11-9 and 12-10 are valid. 11-10 and 12-9 are rejected. Scores above 11 are allowed only after deuce."
    cell.paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph("The final point margin updates skill strength. A close win changes ratings gently; a dominant win makes a larger adjustment.")

    doc.add_heading("5. Rotation and Matchmaking", level=1)
    rotation = [
        ("Winner vs Winner", "Winning partners stay together and wait for another winning team."),
        ("Balanced Remix", "Partners may change. The app balances the selected players using current skill results."),
        ("Pure Random", "Eligible players are reshuffled after every game."),
        ("Rest fairness", "Players who just finished wait behind rested players with fewer completed games."),
    ]
    for label, detail in rotation:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(detail)

    doc.add_heading("6. Team Tournament", level=1)
    doc.add_picture(str(tourney_img), width=Inches(6.45))
    cap = doc.add_paragraph("Figure 3. Round-robin standings and upcoming matches.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor.from_string("71809B")
    doc.add_paragraph("Generate tournament schedule creates every team-versus-team round-robin fixture. The assignment engine prevents one team from playing on two courts simultaneously.")
    doc.add_paragraph("Standings rank teams by wins, then point difference. Entering final scores produces more useful standings than winner-only entry.")

    doc.add_heading("7. Session Data and Offline Use", level=1)
    data_points = [
        "The browser saves every change automatically.",
        "No account or central database is required.",
        "The app keeps working offline after it has been installed or cached.",
        "Export creates a JSON backup; Import restores that backup.",
        "Sessions do not automatically sync between devices.",
        "Clearing browser site data removes the local session.",
    ]
    for item in data_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Organizer Checklist", level=1)
    checks = [
        "Confirm court count and event mode before player check-in.",
        "Test the announcer volume before assigning the first court.",
        "Keep the organizer device charged.",
        "Use final scores when practical; use winner-only entry when courts are moving quickly.",
        "Close courts as reserved time expires.",
        "Export the session before clearing browser data or moving devices.",
    ]
    for item in checks:
        doc.add_paragraph("☐ " + item)

    doc.add_heading("About the Application", level=1)
    doc.add_paragraph("King of Open Play was conceived and directed by Developer King, an engineer in the airline industry who loves data and builds AI-powered applications on the side. AI was used as a development tool for planning, programming, testing, and interface iteration.")
    doc.add_paragraph("The application is designed for free public access as a static Progressive Web App.")

    out = ROOT / "King-of-Open-Play-User-Guide.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_docx())
